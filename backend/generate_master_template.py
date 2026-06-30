import sys
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_real_toc(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_master_template():
    doc = Document()
    
    # Configure alternating headers/footers
    settings = doc.settings.element
    evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
    settings.append(evenAndOddHeaders)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        section.different_first_page_header_footer = True

        # Even Header
        even_header = section.even_page_header
        p_even_head = even_header.paragraphs[0]
        p_even_head.text = "AMCP Format for Formulary Submissions 5.0"
        p_even_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p_even_head.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x14, 0x37, 0x7D) # JMCP blue
            
        # Odd Header
        odd_header = section.header
        p_odd_head = odd_header.paragraphs[0]
        p_odd_head.text = "AMCP Format for Formulary Submissions 5.0"
        p_odd_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p_odd_head.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x14, 0x37, 0x7D) # JMCP blue

        # Even Footer
        even_footer = section.even_page_footer
        p_even_foot = even_footer.paragraphs[0]
        run = p_even_foot.add_run()
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = create_element('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run = p_even_foot.add_run('\tJMCP.org | April 2024 | Vol. 30, No. 4-b')
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        
        pPr = p_even_foot._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), str(int(Inches(6.9) * 1440)))
        tabs.append(tab)
        pPr.append(tabs)

        # Odd Footer
        odd_footer = section.footer
        p_odd_foot = odd_footer.paragraphs[0]
        run = p_odd_foot.add_run('JMCP.org | April 2024 | Vol. 30, No. 4-b\t')
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        
        run = p_odd_foot.add_run()
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = create_element('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        pPr = p_odd_foot._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), str(int(Inches(6.9) * 1440)))
        tabs.append(tab)
        pPr.append(tabs)

    # Modify Normal Style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(10)

    # Cover Page Top Right Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_cit.paragraph_format.space_after = Pt(36)
    r = p_cit.add_run('April 2024\nVolume 30 Number 4-b')
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.font.bold = False
    r.font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(6)
    r = p_title.add_run('AMCP Format for Formulary\nSubmissions 5.0')
    r.font.name = 'Arial'
    r.font.size = Pt(38)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x14, 0x37, 0x7D) # JMCP blue
    
    # website
    p_web = doc.add_paragraph()
    p_web.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_web.paragraph_format.space_after = Pt(24)
    r = p_web.add_run('jmcp.org')
    r.font.name = 'Arial'
    r.font.size = Pt(14.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x14, 0x37, 0x7D)
    
    # Subtitle
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_subtitle.paragraph_format.space_after = Pt(48)
    r = p_subtitle.add_run('Guidance on Submission of Pre-approval and Post-approval Clinical and Economic Information and Evidence')
    r.font.name = 'Arial'
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x87, 0x8A, 0x8F) # Slate grey
    
    # Metadata blocks
    metadata_fields = [
        ('Product:', '[BRAND_NAME] ([GENERIC_NAME])'),
        ('Manufacturer:', '[MANUFACTURER]'),
        ('Version:', '[VERSION_NUMBER]'),
        ('Date:', '[COMPILATION_DATE]')
    ]
    
    for label, marker in metadata_fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(label + ' ')
        r1.font.name = 'Arial'
        r1.font.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0x14, 0x37, 0x7D)
        r2 = p.add_run(marker)
        r2.font.name = 'Arial'
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Copyright footer at bottom
    p_cpy = doc.add_paragraph()
    p_cpy.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cpy.paragraph_format.space_before = Pt(72)
    r = p_cpy.add_run('Copyright © 2024 AMCP')
    r.font.name = 'Arial'
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0x87, 0x8A, 0x8F)

    doc.add_page_break()

    # Table of Contents
    p_toc_head = doc.add_paragraph('Table of Contents')
    p_toc_head.style = doc.styles['Heading 1']
    for r in p_toc_head.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(20)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x14, 0x37, 0x7D)
    
    add_real_toc(doc)
    
    doc.add_page_break()

    # Placeholders for headings and content
    sections_mapping = [
        ('[HEADING_EXEC_SUMMARY]', '[INSERT_EXEC_SUMMARY]'),
        ('[HEADING_PRODUCT_INFO]', '[INSERT_PRODUCT_INFO]'),
        ('[HEADING_CLINICAL_EVIDENCE]', '[INSERT_CLINICAL_EVIDENCE]'),
        ('[HEADING_ECONOMIC_MODEL]', '[INSERT_ECONOMIC_MODEL]'),
        ('[HEADING_SUPPORTING_INFO]', '[INSERT_SUPPORTING_INFO]')
    ]

    for heading_placeholder, content_placeholder in sections_mapping:
        doc.add_paragraph(heading_placeholder)
        doc.add_paragraph(content_placeholder)
        doc.add_page_break()
        
    doc.save('amcp_template.docx')
    print("Successfully generated amcp_template.docx with layout-perfect styles.")

if __name__ == "__main__":
    generate_master_template()
