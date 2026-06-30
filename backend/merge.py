import sys
import os
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docxcompose.composer import Composer

def copy_body_elements(source: Document, destination: Document) -> None:
    elements = list(source.element.body)
    # Remove section properties from the end of the source body if it exists,
    # to avoid messing up the target document page layouts.
    if elements and elements[-1].tag == qn("w:sectPr"):
        elements = elements[:-1]
    for element in elements:
        destination.element.body.append(deepcopy(element))

def merge_with_docxcompose(sources, output_path):
    master = Document(sources[0])
    composer = Composer(master)
    for src in sources[1:]:
        master.add_page_break()
        composer.append(Document(src))
    composer.save(output_path)

def merge_manual(sources, output_path):
    merged = Document(sources[0])
    for src in sources[1:]:
        merged.add_page_break()
        copy_body_elements(Document(src), merged)
    merged.save(output_path)

# ---------------------------------------------------------
# Extensible Conversion Layer
# ---------------------------------------------------------

def convert_pdf_to_docx(input_path: Path) -> Path:
    from pdf2docx import Converter
    temp_docx_path = input_path.with_suffix(".temp.docx")
    print(f"Converting PDF: {input_path.name} -> {temp_docx_path.name}")
    cv = Converter(str(input_path))
    cv.convert(str(temp_docx_path), multi_processing=True)
    cv.close()
    return temp_docx_path

def convert_ppt_to_docx(input_path: Path, temp_files: list) -> Path:
    import win32com.client
    temp_pdf_path = input_path.with_suffix(".temp.pdf")
    temp_docx_path = input_path.with_suffix(".temp.docx")
    
    # Track the PDF path immediately for cleanup in case of crash later
    temp_files.append(temp_pdf_path)
    
    print(f"Converting PowerPoint: {input_path.name} -> PDF -> {temp_docx_path.name}")
    
    abs_input = str(input_path.resolve())
    abs_pdf = str(temp_pdf_path.resolve())
    
    powerpoint = win32com.client.Dispatch("Powerpoint.Application")
    try:
        # Open presentation in background (WithWindow=False)
        presentation = powerpoint.Presentations.Open(abs_input, WithWindow=False)
        # Format 32 is ppSaveAsPDF
        presentation.SaveAs(abs_pdf, 32)
        presentation.Close()
    finally:
        try:
            powerpoint.Quit()
        except Exception:
            pass
            
    # Now convert the temp PDF to DOCX using pdf2docx
    from pdf2docx import Converter
    cv = Converter(abs_pdf)
    cv.convert(str(temp_docx_path), multi_processing=True)
    cv.close()
    
    return temp_docx_path

def convert_xlsx_to_docx(input_path: Path) -> Path:
    raise NotImplementedError("XLSX conversion is designed but not yet active.")

def convert_txt_to_docx(input_path: Path) -> Path:
    raise NotImplementedError("TXT conversion is designed but not yet active.")

def convert_rtf_to_docx(input_path: Path) -> Path:
    raise NotImplementedError("RTF conversion is designed but not yet active.")

def convert_odt_to_docx(input_path: Path) -> Path:
    raise NotImplementedError("ODT conversion is designed but not yet active.")

# Mappings of extensions to their converter functions
CONVERTER_REGISTRY = {
    ".pdf": convert_pdf_to_docx,
    ".ppt": convert_ppt_to_docx,
    ".pptx": convert_ppt_to_docx,
    ".xlsx": convert_xlsx_to_docx,
    ".txt": convert_txt_to_docx,
    ".rtf": convert_rtf_to_docx,
    ".odt": convert_odt_to_docx,
}

def resolve_and_convert(file_path: Path, temp_files: list) -> Path:
    ext = file_path.suffix.lower()
    if ext == ".docx":
        return file_path
    
    if ext in CONVERTER_REGISTRY:
        try:
            if ext in (".ppt", ".pptx"):
                converted_path = CONVERTER_REGISTRY[ext](file_path, temp_files)
            else:
                converted_path = CONVERTER_REGISTRY[ext](file_path)
            temp_files.append(converted_path)
            return converted_path
        except NotImplementedError as nie:
            raise ValueError(f"Support for converting {ext.upper()} files is designed but not yet active: {nie}")
        except Exception as e:
            raise RuntimeError(f"Failed to convert {file_path.name} to DOCX: {e}")
    else:
        raise ValueError(f"Unsupported file format for conversion: {ext}")

# ---------------------------------------------------------
# Main Flow
# ---------------------------------------------------------

def main():
    if len(sys.argv) < 4:
        print("Usage: python merge.py <output_path> <input_path_1> <input_path_2> [input_path_3 ...]")
        sys.exit(1)

    output_path = Path(sys.argv[1])
    raw_sources = [Path(p) for p in sys.argv[2:]]

    temp_files = []
    processed_sources = []

    print(f"Processing and merging {len(raw_sources)} files into: {output_path}")
    
    try:
        # Convert non-docx files using the converter registry
        for idx, src in enumerate(raw_sources):
            if not src.exists():
                raise FileNotFoundError(f"Input file not found: {src}")
            docx_path = resolve_and_convert(src, temp_files)
            processed_sources.append(docx_path)
        
        # Merge all resulting DOCX files
        print(f"Merging intermediate DOCX files:")
        for idx, src in enumerate(processed_sources):
            print(f"  [{idx + 1}] {src}")

        try:
            merge_with_docxcompose(processed_sources, output_path)
            print("DOCX merged successfully with docxcompose:", output_path)
        except Exception as exc:
            print("docxcompose failed, using manual fallback:", exc)
            try:
                merge_manual(processed_sources, output_path)
                print("DOCX merged successfully with fallback:", output_path)
            except Exception as fallback_exc:
                print("Fallback merge also failed:", fallback_exc)
                sys.exit(1)

    except Exception as process_exc:
        print(f"Error during document processing: {process_exc}", file=sys.stderr)
        sys.exit(1)

    finally:
        # Clean up temporary DOCX files
        if temp_files:
            print("Cleaning up temporary converted files:")
            for temp_f in temp_files:
                try:
                    if temp_f.exists():
                        print(f"  Removing: {temp_f.name}")
                        os.remove(temp_f)
                except Exception as cleanup_err:
                    print(f"  Failed to remove temporary file {temp_f.name}: {cleanup_err}", file=sys.stderr)

if __name__ == "__main__":
    main()
