import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Colors
COLOR_JMCP_BLUE = RGBColor(0x14, 0x37, 0x7D)    # #14377D
COLOR_SLATE_GRAY = RGBColor(0x87, 0x8A, 0x8F)   # #878A8F
COLOR_CHARCOAL = RGBColor(0x47, 0x4C, 0x55)     # #474C55
COLOR_PURPLE = RGBColor(0x6B, 0x14, 0x6A)       # #6B146A
HEX_JMCP_BLUE = "14377D"
HEX_SLATE_GRAY = "878A8F"
HEX_CHARCOAL = "474C55"
HEX_PURPLE = "6B146A"
HEX_VERY_LIGHT_GRAY = "E4E5E6"

# Helper: Create OxmlElement
def create_element(name):
    return OxmlElement(name)

# Helper: Create attribute
def create_attribute(element, name, value):
    element.set(qn(name), value)

# Helper: Clear table borders
def clear_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# Helper: Set cell margins (padding) in dxa (1 inch = 1440 dxa)
def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m_el = OxmlElement(f'w:{margin}')
        m_el.set(qn('w:w'), str(val))
        m_el.set(qn('w:type'), 'dxa')
        tcMar.append(m_el)
    tcPr.append(tcMar)

# Helper: Set cell shading fill color
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# Helper: Set cell width in inches
def set_cell_width(cell, width_in):
    cell.width = Inches(width_in)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_in * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# Helper: Set cell vertical alignment
def set_cell_vertical_alignment(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

# Helper: Set cell text direction (btLr for bottom-to-top vertical text)
def set_cell_text_direction(cell, direction="btLr"):
    tcPr = cell._tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)
    tcPr.append(textDirection)

# 1. Apply Global Theme to Document Styles
def apply_global_theme(doc):
    # Configure default styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Lora'
    style_normal.font.size = Pt(9.0)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)
    
    # Configure Heading 1
    try:
        style_h1 = doc.styles['Heading 1']
        style_h1.font.name = 'Asap'
        style_h1.font.size = Pt(16)
        style_h1.font.bold = True
        style_h1.font.color.rgb = COLOR_PURPLE
        style_h1.paragraph_format.space_before = Pt(12)
        style_h1.paragraph_format.space_after = Pt(6)
    except KeyError:
        pass
        
    # Configure Heading 2
    try:
        style_h2 = doc.styles['Heading 2']
        style_h2.font.name = 'Asap'
        style_h2.font.size = Pt(12)
        style_h2.font.bold = True
        style_h2.font.color.rgb = COLOR_JMCP_BLUE
        style_h2.paragraph_format.space_before = Pt(10)
        style_h2.paragraph_format.space_after = Pt(4)
    except KeyError:
        pass

    # Configure Heading 3
    try:
        style_h3 = doc.styles['Heading 3']
        style_h3.font.name = 'Asap'
        style_h3.font.size = Pt(10)
        style_h3.font.bold = True
        style_h3.font.color.rgb = COLOR_JMCP_BLUE
        style_h3.paragraph_format.space_before = Pt(8)
        style_h3.paragraph_format.space_after = Pt(2)
    except KeyError:
        pass

# 2. Configure Section Layout
def configure_section_layout(section, layout_type="content"):
    # Page size
    section.page_width = Inches(8.25)
    section.page_height = Inches(10.875)
    
    if layout_type == "cover":
        # Full-bleed margins
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.different_first_page_header_footer = True
        
    elif layout_type == "toc":
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.different_first_page_header_footer = False
        
        # Single column
        sectPr = section._sectPr
        cols = sectPr.xpath('w:cols')
        if cols:
            cols[0].set(qn('w:num'), '1')
            
    elif layout_type == "content":
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.different_first_page_header_footer = False
        
        # Two columns with 0.25 inch spacing
        sectPr = section._sectPr
        cols = sectPr.xpath('w:cols')
        if cols:
            cols[0].set(qn('w:num'), '2')
            cols[0].set(qn('w:space'), '360')  # 0.25 inches = 360 dxa
        else:
            col_el = OxmlElement('w:cols')
            col_el.set(qn('w:num'), '2')
            col_el.set(qn('w:space'), '360')
            sectPr.append(col_el)

# 3. Setup Alternating Header and Footer on Section
def setup_headers_footers(section):
    # Setup Even Header
    even_header = section.even_page_header
    p_even = even_header.paragraphs[0]
    p_even.text = ""
    # Header table: 1 row, 2 cells
    tbl_even = even_header.add_table(1, 2, Inches(7.25))
    tbl_even.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(tbl_even)
    row_even = tbl_even.rows[0]
    row_even.height = Inches(0.3)
    
    # Cell 1 (Page Number): width 0.35 in, purple, white vertical page number
    cell_p = row_even.cells[0]
    set_cell_width(cell_p, 0.35)
    set_cell_shading(cell_p, HEX_PURPLE)
    set_cell_vertical_alignment(cell_p, "center")
    set_cell_margins(cell_p, top=40, bottom=40, left=40, right=40)
    p_num = cell_p.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run()
    r_num.font.name = 'Asap'
    r_num.font.size = Pt(9)
    r_num.font.bold = True
    r_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    fld1 = create_element('w:fldChar')
    create_attribute(fld1, 'w:fldCharType', 'begin')
    instr = create_element('w:instrText')
    instr.text = "PAGE"
    fld2 = create_element('w:fldChar')
    create_attribute(fld2, 'w:fldCharType', 'end')
    r_num._r.append(fld1)
    r_num._r.append(instr)
    r_num._r.append(fld2)
    
    # Cell 2 (Title): width 6.90 in, Charcoal text
    cell_t = row_even.cells[1]
    set_cell_width(cell_t, 6.90)
    set_cell_vertical_alignment(cell_t, "center")
    set_cell_margins(cell_t, top=40, bottom=40, left=144, right=40) # left indent in cell padding (144 dxa = 0.1 inch)
    p_title = cell_t.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run("AMCP FORMAT 5.0")
    r_title.font.name = 'Asap'
    r_title.font.size = Pt(9)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_CHARCOAL
    
    # Setup Odd Header
    odd_header = section.header
    p_odd = odd_header.paragraphs[0]
    p_odd.text = ""
    tbl_odd = odd_header.add_table(1, 2, Inches(7.25))
    tbl_odd.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(tbl_odd)
    row_odd = tbl_odd.rows[0]
    row_odd.height = Inches(0.3)
    
    # Cell 1 (Title): width 6.90 in, right aligned
    cell_ot = row_odd.cells[0]
    set_cell_width(cell_ot, 6.90)
    set_cell_vertical_alignment(cell_ot, "center")
    set_cell_margins(cell_ot, top=40, bottom=40, left=40, right=144)
    p_ot = cell_ot.paragraphs[0]
    p_ot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ot = p_ot.add_run("AMCP FORMAT 5.0")
    r_ot.font.name = 'Asap'
    r_ot.font.size = Pt(9)
    r_ot.font.bold = True
    r_ot.font.color.rgb = COLOR_CHARCOAL
    
    # Cell 2 (Page Number): width 0.35 in, purple, white text
    cell_op = row_odd.cells[1]
    set_cell_width(cell_op, 0.35)
    set_cell_shading(cell_op, HEX_PURPLE)
    set_cell_vertical_alignment(cell_op, "center")
    set_cell_margins(cell_op, top=40, bottom=40, left=40, right=40)
    p_onum = cell_op.paragraphs[0]
    p_onum.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_onum = p_onum.add_run()
    r_onum.font.name = 'Asap'
    r_onum.font.size = Pt(9)
    r_onum.font.bold = True
    r_onum.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    fld1 = create_element('w:fldChar')
    create_attribute(fld1, 'w:fldCharType', 'begin')
    instr = create_element('w:instrText')
    instr.text = "PAGE"
    fld2 = create_element('w:fldChar')
    create_attribute(fld2, 'w:fldCharType', 'end')
    r_onum._r.append(fld1)
    r_onum._r.append(instr)
    r_onum._r.append(fld2)

    # Setup Even Footer: JMCP.org | April 2024 | Vol. 30, No. 4-B
    even_footer = section.even_page_footer
    p_ef = even_footer.paragraphs[0]
    p_ef.text = ""
    p_ef.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p_ef.add_run("JMCP")
    r1.font.name = 'Asap'
    r1.font.size = Pt(7)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_JMCP_BLUE
    
    r2 = p_ef.add_run(".org")
    r2.font.name = 'Asap'
    r2.font.size = Pt(7)
    r2.font.color.rgb = COLOR_JMCP_BLUE
    
    r3 = p_ef.add_run(" | April 2024 | Vol. 30, No. 4-B")
    r3.font.name = 'Asap'
    r3.font.size = Pt(7)
    r3.font.color.rgb = COLOR_CHARCOAL

    # Setup Odd Footer: Vol. 30, No. 4-b | April 2024 | JMCP.org
    odd_footer = section.footer
    p_of = odd_footer.paragraphs[0]
    p_of.text = ""
    p_of.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_of.add_run("Vol. 30, No. 4-b | April 2024 | ")
    r1.font.name = 'Asap'
    r1.font.size = Pt(7)
    r1.font.color.rgb = COLOR_CHARCOAL
    
    r2 = p_of.add_run("JMCP")
    r2.font.name = 'Asap'
    r2.font.size = Pt(7)
    r2.font.bold = True
    r2.font.color.rgb = COLOR_JMCP_BLUE
    
    r3 = p_of.add_run(".org")
    r3.font.name = 'Asap'
    r3.font.size = Pt(7)
    r3.font.color.rgb = COLOR_JMCP_BLUE

# 4. Create Cover Page Component
def create_cover_page_component(doc):
    # The document must have its first section set to "cover" layout
    section = doc.sections[0]
    configure_section_layout(section, "cover")
    
    # Add full-page table
    table = doc.add_table(1, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(table)
    
    row = table.rows[0]
    # Set total height to full page height (10.875 inches)
    row.height = Inches(10.875)
    
    # Cell 1: Vertical Gray Bar on left edge
    cell_bar = row.cells[0]
    set_cell_width(cell_bar, 0.25)
    set_cell_shading(cell_bar, HEX_SLATE_GRAY)
    set_cell_vertical_alignment(cell_bar, "center")
    set_cell_text_direction(cell_bar, "btLr")
    set_cell_margins(cell_bar, top=0, bottom=0, left=0, right=0)
    
    p_bar = cell_bar.paragraphs[0]
    p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bar = p_bar.add_run("2024        30  4B         JMCP")
    r_bar.font.name = 'Asap'
    r_bar.font.size = Pt(7.5)
    r_bar.font.bold = True
    r_bar.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Cell 2: Content Block
    cell_content = row.cells[1]
    set_cell_width(cell_content, 8.00)
    set_cell_margins(cell_content, top=1152, bottom=1152, left=720, right=720) # top/bottom padding 0.8", left/right padding 0.5"
    set_cell_vertical_alignment(cell_content, "top")
    
    # Inside Cell 2: Top Logo & Citation Nested Table
    p_start = cell_content.paragraphs[0]
    p_start.text = "" # Clear default
    
    logo_table = cell_content.add_table(1, 2)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(logo_table)
    logo_row = logo_table.rows[0]
    logo_row.height = Inches(1.15)
    
    # Citation (Left cell in logo table)
    cell_cit = logo_row.cells[0]
    set_cell_width(cell_cit, 5.75)
    set_cell_vertical_alignment(cell_cit, "center")
    set_cell_margins(cell_cit, top=0, bottom=0, left=0, right=0)
    p_cit = cell_cit.paragraphs[0]
    p_cit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_cit = p_cit.add_run("April 2024\nVolume 30 Number 4-b")
    r_cit.font.name = 'Asap'
    r_cit.font.size = Pt(9)
    r_cit.font.color.rgb = RGBColor(0, 0, 0)
    
    # Purple Logo (Right cell in logo table)
    cell_logo = logo_row.cells[1]
    set_cell_width(cell_logo, 1.25)
    set_cell_shading(cell_logo, HEX_PURPLE)
    set_cell_vertical_alignment(cell_logo, "center")
    set_cell_margins(cell_logo, top=40, bottom=40, left=40, right=40)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("JMCP")
    r_logo.font.name = 'Asap'
    r_logo.font.size = Pt(18)
    r_logo.font.bold = True
    r_logo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Spacer
    p_space1 = cell_content.add_paragraph()
    p_space1.paragraph_format.space_before = Pt(108) # Spacing before Title
    
    # Title
    p_title = cell_content.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(6)
    r_t = p_title.add_run("AMCP Format for Formulary\nSubmissions 5.0")
    r_t.font.name = 'Asap'
    r_t.font.size = Pt(38)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_JMCP_BLUE
    
    # Website
    p_web = cell_content.add_paragraph()
    p_web.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_web.paragraph_format.space_after = Pt(24)
    r_w = p_web.add_run("jmcp.org")
    r_w.font.name = 'Asap'
    r_w.font.size = Pt(14.5)
    r_w.font.italic = True
    r_w.font.color.rgb = COLOR_JMCP_BLUE
    
    # Subtitle
    p_subtitle = cell_content.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_subtitle.paragraph_format.space_after = Pt(48)
    r_sub = p_subtitle.add_run("Guidance on Submission of Pre-approval and Post-approval Clinical and Economic Information and Evidence")
    r_sub.font.name = 'Asap'
    r_sub.font.size = Pt(22)
    r_sub.font.color.rgb = COLOR_SLATE_GRAY
    
    # Metadata nested table
    meta_table = cell_content.add_table(4, 2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    clear_table_borders(meta_table)
    
    metadata_fields = [
        ("Product:", "[BRAND_NAME] ([GENERIC_NAME])"),
        ("Manufacturer:", "[MANUFACTURER]"),
        ("Version:", "[VERSION_NUMBER]"),
        ("Date:", "[COMPILATION_DATE]")
    ]
    
    for idx, (label, value) in enumerate(metadata_fields):
        meta_row = meta_table.rows[idx]
        meta_row.height = Inches(0.25)
        
        # Label cell
        cell_l = meta_row.cells[0]
        set_cell_width(cell_l, 1.8)
        p_l = cell_l.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_l = p_l.add_run(label)
        r_l.font.name = 'Asap'
        r_l.font.bold = True
        r_l.font.size = Pt(12)
        r_l.font.color.rgb = COLOR_JMCP_BLUE
        
        # Value cell
        cell_v = meta_row.cells[1]
        set_cell_width(cell_v, 5.2)
        p_v = cell_v.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_v = p_v.add_run(value)
        r_v.font.name = 'Asap'
        r_v.font.size = Pt(12)
        r_v.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    # Spacer before copyright
    p_space2 = cell_content.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(72)
    
    # Copyright
    p_cpy = cell_content.add_paragraph()
    p_cpy.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cpy = p_cpy.add_run("Copyright © 2024 AMCP")
    r_cpy.font.name = 'Asap'
    r_cpy.font.size = Pt(7.5)
    r_cpy.font.color.rgb = COLOR_SLATE_GRAY

# 5. Create Table of Contents Component
def create_toc_component(doc):
    # Add section break for TOC
    new_sect = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    configure_section_layout(new_sect, "toc")
    setup_headers_footers(new_sect)
    
    # Heading
    p_toc_head = doc.add_paragraph('Table of Contents')
    p_toc_head.style = doc.styles['Heading 1']
    for r in p_toc_head.runs:
        r.font.name = 'Asap'
        r.font.size = Pt(20)
        r.font.bold = True
        r.font.color.rgb = COLOR_JMCP_BLUE
        
    # TOC field codes
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# 6. Format Section Heading (reusable component)
def create_section_heading_component(paragraph, num, text):
    try:
        paragraph.style = paragraph.part.document.styles['Heading 1']
    except Exception:
        pass
    paragraph.text = ""
    
    cyan_run = paragraph.add_run(f"SECTION {num} \n")
    cyan_run.font.name = 'Asap'
    cyan_run.font.size = Pt(10)
    cyan_run.font.bold = True
    cyan_run.font.color.rgb = RGBColor(0, 126, 197) # #007EC5
    
    plum_run = paragraph.add_run(text)
    plum_run.font.name = 'Asap'
    plum_run.font.size = Pt(16)
    plum_run.font.bold = True
    plum_run.font.color.rgb = RGBColor(108, 16, 91)  # #6C105B
    
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
