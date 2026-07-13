import sys
import os
import json
import re
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
import template_components

# Helper: Clean XML incompatible characters to prevent ValueError in lxml
def clean_xml_compatible_text(text):
    if not isinstance(text, str):
        return ""
    cleaned = []
    for char in text:
        cp = ord(char)
        if (cp == 0x9 or cp == 0xA or cp == 0xD or
            (0x20 <= cp <= 0xD7FF) or
            (0xE000 <= cp <= 0xFFFD) or
            (0x10000 <= cp <= 0x10FFFF)):
            cleaned.append(char)
    return "".join(cleaned)

# Layout and styling helper functions have been moved to template_components.py to centralize presentation details.

def convert_pdf_to_docx(pdf_path, docx_path):
    print(f"Converting PDF {pdf_path} to DOCX...")
    cv = Converter(pdf_path)
    cv.convert(docx_path)
    cv.close()
    return docx_path

def replace_text_in_document(doc, old_text, new_text):
    new_text = clean_xml_compatible_text(new_text)
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(old_text, new_text)
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        inline = p.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                inline[i].text = inline[i].text.replace(old_text, new_text)
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

# Iterate paragraphs and tables sequentially
def iter_block_items(doc):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('p'):
            yield docx.text.paragraph.Paragraph(child, doc)
        elif child.tag.endswith('tbl'):
            yield docx.table.Table(child, doc)

# Parse drawings inside paragraph xml and write raw bytes
def extract_runs_images(p, temp_dir):
    image_paths = []
    try:
        blips = p._element.xpath('.//*[local-name()="blip"]')
        for blip in blips:
            embed_attr = None
            for attr in blip.attrib:
                if attr.endswith('embed'):
                    embed_attr = attr
                    break
            if embed_attr:
                r_id = blip.attrib[embed_attr]
                if r_id in p.part.related_parts:
                    part = p.part.related_parts[r_id]
                    if "image" in part.content_type:
                        ext = part.content_type.split('/')[-1]
                        filename = f"img_{r_id}_{len(image_paths)}.{ext}"
                        path = os.path.join(temp_dir, filename)
                        with open(path, 'wb') as img_f:
                            img_f.write(part.blob)
                        image_paths.append(path)
    except Exception as e:
        print(f"Error extracting image from paragraph: {e}")
    return image_paths

# Extract structured content from DOCX
def extract_docx_content_model(source_docx_path, temp_dir):
    source_doc = Document(source_docx_path)
    blocks = list(iter_block_items(source_doc))
    
    header_footer_patterns = [
        re.compile(r"^Page\s+\d+", re.IGNORECASE),
        re.compile(r"^\d+\s+of\s+\d+", re.IGNORECASE),
        re.compile(r"^\d+\s*/\s*\d+", re.IGNORECASE),
        re.compile(r"JMCP\.org", re.IGNORECASE),
        re.compile(r"Volume\s+\d+", re.IGNORECASE),
        re.compile(r"Number\s+\d+\S*", re.IGNORECASE),
        re.compile(r"\.{4,}", re.IGNORECASE),
    ]
    
    model_blocks = []
    for b in blocks:
        if isinstance(b, docx.text.paragraph.Paragraph):
            text = b.text.strip()
            # Filter margin/navigation tags
            skip = False
            for cp in header_footer_patterns:
                if cp.search(text):
                    skip = True
                    break
            if skip:
                continue
                
            # Check for images inside paragraph xml
            img_paths = extract_runs_images(b, temp_dir)
            if img_paths:
                for path in img_paths:
                    model_blocks.append({
                        'type': 'image',
                        'path': path,
                        'caption': text if text else "Figure"
                    })
                continue
                
            if not text:
                continue
                
            # Classify Heading Level based on style name
            style_name = b.style.name.lower() if b.style else ""
            heading_level = 0
            if 'heading' in style_name:
                if '1' in style_name: heading_level = 1
                elif '2' in style_name: heading_level = 2
                elif '3' in style_name: heading_level = 3
                else: heading_level = 2
                
            # Normalize text runs
            runs_model = []
            for r in b.runs:
                if r.text:
                    runs_model.append({
                        'text': r.text,
                        'bold': bool(r.bold),
                        'italic': bool(r.italic)
                    })
                    
            model_blocks.append({
                'type': 'paragraph',
                'text': text,
                'heading_level': heading_level,
                'runs': runs_model
            })
            
        elif isinstance(b, docx.table.Table):
            rows_data = []
            for r in b.rows:
                cell_texts = [cell.text.strip() for cell in r.cells]
                rows_data.append(cell_texts)
            if rows_data:
                model_blocks.append({
                    'type': 'table',
                    'rows': rows_data
                })
                
    return model_blocks

# Extract structured content from PDF Page
def extract_pdf_page_blocks(page, temp_dir):
    blocks = []
    
    # 1. Extract tables first
    tables = page.find_tables()
    table_bboxes = []
    for t in tables:
        grid = t.extract()
        table_bboxes.append((t.bbox, grid))
        
    # 2. Extract text blocks
    text_blocks = page.get_text("blocks")
    
    # Filter margin/header/footer text out based on y0/y1 thresholds
    page_height = page.rect.height
    top_limit = 65
    bottom_limit = page_height - 65
    
    for tb in text_blocks:
        x0, y0, x1, y1, text, block_no, block_type = tb
        
        # Filter page borders/navigation headers/footers
        if y0 < top_limit or y1 > bottom_limit:
            continue
            
        # Ignore text blocks falling inside table boundaries
        is_inside_table = False
        for bbox, grid in table_bboxes:
            tx0, ty0, tx1, ty1 = bbox
            if ty0 - 2 <= y0 <= ty1 + 2:
                is_inside_table = True
                break
        if is_inside_table:
            continue
            
        text_clean = text.strip()
        if not text_clean:
            continue
            
        # Check if text is a heading or bullet list
        heading_level = 0
        if len(text_clean) < 80 and (text_clean.isupper() or re.match(r'^\d+\.\d+', text_clean)):
            heading_level = 2
            
        blocks.append({
            'type': 'paragraph',
            'y0': y0,
            'text': text_clean,
            'heading_level': heading_level,
            'runs': [{'text': text_clean, 'bold': False, 'italic': False}]
        })
        
    # 3. Add tables
    for bbox, grid in table_bboxes:
        tx0, ty0, tx1, ty1 = bbox
        blocks.append({
            'type': 'table',
            'y0': ty0,
            'rows': grid
        })
        
    # 4. Extract images
    image_list = page.get_images(full=True)
    for img_idx, img in enumerate(image_list):
        xref = img[0]
        rects = page.get_image_rects(xref)
        if rects:
            rect = rects[0]
            y0 = rect.y0
            
            try:
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image["image"]
                ext = base_image["ext"]
                filename = f"pdf_img_{xref}_{img_idx}.{ext}"
                img_path = os.path.join(temp_dir, filename)
                with open(img_path, 'wb') as img_f:
                    img_f.write(image_bytes)
                    
                blocks.append({
                    'type': 'image',
                    'y0': y0,
                    'path': img_path,
                    'caption': "Figure"
                })
            except Exception as e:
                print(f"Error extracting PDF image xref {xref}: {e}")
                
    # Sort blocks page-order vertically
    blocks.sort(key=lambda b: b['y0'])
    return blocks

# Prune content blocks for target section
def prune_blocks_list(blocks, section_key):
    patterns = {
        'exec_summary': [r'1\.0\s+EXECUTIVE\s+SUMMARY', r'1\.0B\s+EXECUTIVE\s+SUMMARY', r'1\.0A\s+EXECUTIVE\s+SUMMARY'],
        'highlights': [r'1\.0\s+HIGHLIGHTS', r'1\.0A\s+HIGHLIGHTS', r'1\.0C\s+HIGHLIGHTS'],
        'product_info': [r'2\.0\s+PRODUCT\s+INFORMATION', r'2\.0B\s+PRODUCT\s+INFORMATION', r'2\.0A\s+PRODUCT\s+INFORMATION', r'2\.0\s+DISEASE\s+DESCRIPTION', r'2\.0B\s+DISEASE\s+DESCRIPTION'],
        'clinical': [r'3\.0\s+CLINICAL\s+EVIDENCE', r'3\.0B\s+CLINICAL\s+EVIDENCE', r'3\.0A\s+CLINICAL\s+EVIDENCE'],
        'economic': [r'4\.0\s+ECONOMIC\s+VALUE', r'4\.0B\s+ECONOMIC\s+VALUE', r'4\.0A\s+ECONOMIC\s+VALUE', r'4\.0\s+ECONOMIC\s+INFORMATION', r'4\.0A\s+ECONOMIC\s+INFORMATION'],
        'supporting': [r'5\.0\s+ADDITIONAL\s+SUPPORTING', r'5\.0B\s+ADDITIONAL\s+SUPPORTING', r'5\.0A\s+ADDITIONAL\s+SUPPORTING'],
        'references': [r'6\.0\s+DOSSIER\s+APPENDICES', r'6\.0B\s+DOSSIER\s+APPENDICES', r'6\.0A\s+DOSSIER\s+APPENDICES', r'6\.0\s+Dossier\s+Appendices', r'References', r'REFERENCES']
    }
    
    all_start_patterns = []
    for k, pats in patterns.items():
        all_start_patterns.extend(pats)
        
    target_pats = patterns.get(section_key, [])
    compiled_targets = [re.compile(p, re.IGNORECASE) for p in target_pats]
    compiled_stops = [re.compile(p, re.IGNORECASE) for p in all_start_patterns if p not in target_pats]
    
    start_idx = -1
    end_idx = -1
    
    for idx, b in enumerate(blocks):
        if b['type'] == 'paragraph':
            text = b.get('text', '').strip()
            if not text:
                continue
            
            if start_idx == -1:
                for cp in compiled_targets:
                    if cp.search(text):
                        start_idx = idx
                        break
            elif end_idx == -1:
                for cp in compiled_stops:
                    if cp.search(text):
                        end_idx = idx
                        break
                        
    if start_idx == -1:
        start_idx = 0
        end_idx = len(blocks)
    else:
        # Skip heading itself
        start_idx += 1
        if end_idx == -1:
            end_idx = len(blocks)
            
    return blocks[start_idx:end_idx]

# Normalizes OCR hyphenation and text spacing
def normalize_text_runs(text):
    # Remove hyphens split across lines
    text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
    # Convert single newlines to spaces
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
    # Remove extra spaces
    text = re.sub(r' +', ' ', text)
    return text.strip()

# Unified source content extractor
def extract_section_blocks(file_path, section_key, temp_dir):
    suffix = os.path.splitext(file_path)[1].lower()
    
    raw_blocks = []
    if suffix == '.pdf':
        print(f"Extracting PDF content: {file_path}")
        try:
            pdf_doc = fitz.open(file_path)
            for page in pdf_doc:
                raw_blocks.extend(extract_pdf_page_blocks(page, temp_dir))
            pdf_doc.close()
        except Exception as e:
            print(f"Error parsing PDF with PyMuPDF: {e}")
            
    elif suffix == '.docx':
        print(f"Extracting DOCX content: {file_path}")
        try:
            raw_blocks = extract_docx_content_model(file_path, temp_dir)
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            
    # Prune down to section content range
    section_blocks = prune_blocks_list(raw_blocks, section_key)
    
    # Normalize text in paragraphs
    for b in section_blocks:
        if b['type'] == 'paragraph':
            b['text'] = normalize_text_runs(b.get('text', ''))
            for r in b.get('runs', []):
                r['text'] = normalize_text_runs(r.get('text', ''))
                
    return section_blocks

# Render structured blocks directly before placeholder
def render_content_blocks(doc, placeholder_paragraph, blocks):
    p_element = placeholder_paragraph._element
    body_element = p_element.getparent()
    
    for b in blocks:
        if b['type'] == 'paragraph':
            text = b.get('text', '').strip()
            if not text:
                continue
                
            p = doc.add_paragraph()
            heading_lvl = b.get('heading_level', 0)
            template_components.format_paragraph_styling(p, heading_lvl)
                
            for r_data in b.get('runs', []):
                r = p.add_run(clean_xml_compatible_text(r_data['text']))
                r.bold = r_data.get('bold', False)
                r.italic = r_data.get('italic', False)
                
            body_element.insert(body_element.index(p_element), p._element)
            
        elif b['type'] == 'table':
            rows_data = b.get('rows', [])
            if not rows_data or not rows_data[0]:
                continue
                
            row_count = len(rows_data)
            col_count = len(rows_data[0])
            
            table = doc.add_table(row_count, col_count)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            template_components.clear_table_borders(table)
            
            for r_idx, row_cells in enumerate(rows_data):
                row = table.rows[r_idx]
                template_components.make_row_cant_split(row)
                if r_idx == 0:
                    template_components.make_row_tbl_header(row)
                    
                for c_idx, cell_value in enumerate(row_cells):
                    cell = row.cells[c_idx]
                    
                    # Compute shading color for alternating rows
                    shading_color = None
                    if r_idx > 0:
                        shading_color = 'F3F4F6' if r_idx % 2 == 1 else 'FFFFFF'
                        
                    # Delegate styling to template_components cell formatter
                    template_components.format_table_cell(
                        cell=cell,
                        text=clean_xml_compatible_text(str(cell_value)),
                        is_header=(r_idx == 0),
                        width_in=3.5 / col_count,
                        shading_color=shading_color
                    )
                                
            body_element.insert(body_element.index(p_element), table._element)
            
        elif b['type'] == 'image':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True # Anchor figure with caption
            
            run = p.add_run()
            try:
                run.add_picture(b['path'], width=Inches(3.5))
            except Exception as e:
                print(f"Error rendering image block: {e}")
                
            body_element.insert(body_element.index(p_element), p._element)
            
            p_cap = doc.add_paragraph()
            template_components.format_figure_caption(p_cap, clean_xml_compatible_text(f"Figure: {b['caption']}"))
            
            body_element.insert(body_element.index(p_element), p_cap._element)

    body_element.remove(p_element)

# Enable automatic TOC updates on open
def enable_update_fields_on_open(doc):
    try:
        doc.settings.update_fields = True
    except Exception:
        pass
    
    try:
        settings_elm = doc.settings.element
        uf = settings_elm.find(qn('w:updateFields'))
        if uf is None:
            uf = OxmlElement('w:updateFields')
            uf.set(qn('w:val'), 'true')
            settings_elm.append(uf)
    except Exception as e:
        print(f"Error setting update fields flag: {e}")

def compile_dossier(config_path, output_path):
    with open(config_path, 'r') as f:
        config = json.load(f)
        
    metadata = config.get('metadata', {})
    sections = config.get('sections', {})
    dossier_type = config.get('dossierType', 'format-a')

    # Setup temporary directory for images/media extraction
    temp_dir = Path('./temp_extract')
    if temp_dir.exists():
        shutil.rmtree(temp_dir)
    temp_dir.mkdir(parents=True, exist_ok=True)

    templates_dir = Path(__file__).parent / 'templates'
    
    # 1. Load correct high-fidelity master template (Format A or Format B)
    template_filename = 'amcp_template_format_b.docx' if dossier_type == 'format-b' else 'amcp_template_format_a.docx'
    template_path = templates_dir / template_filename
    
    if not template_path.exists():
        raise FileNotFoundError(f"Unified master template not found at {template_path}. Please run generate_templates.py first.")
        
    print(f"Loading Master Template: {template_filename}...")
    doc = Document(str(template_path))
    
    # 2. Inject cover page metadata
    replace_text_in_document(doc, '[BRAND_NAME]', metadata.get('brandName', ''))
    replace_text_in_document(doc, '[GENERIC_NAME]', metadata.get('genericName', ''))
    replace_text_in_document(doc, '[MANUFACTURER]', metadata.get('manufacturer', ''))
    replace_text_in_document(doc, '[VERSION_NUMBER]', metadata.get('versionNumber', ''))
    replace_text_in_document(doc, '[COMPILATION_DATE]', metadata.get('compilationDate', ''))
    
    # 3. Process sections
    section_keys = ['highlights' if dossier_type == 'format-a' else 'exec_summary', 'product_info', 'clinical', 'economic', 'supporting', 'references']
    
    for key in section_keys:
        placeholder_text = f"[INSERT_{key.upper()}]"
        
        # Locate placeholder paragraph in template
        placeholder_p = None
        for p in doc.paragraphs:
            if placeholder_text in p.text:
                placeholder_p = p
                break
                
        if not placeholder_p:
            continue
            
        file_path = sections.get(key)
        if file_path and os.path.exists(file_path):
            # Extract structured content blocks using our deterministic pipeline
            blocks = extract_section_blocks(file_path, key, str(temp_dir))
            
            if blocks:
                print(f"Rendering {len(blocks)} content blocks into placeholder {placeholder_text}...")
                render_content_blocks(doc, placeholder_p, blocks)
            else:
                placeholder_p.text = "No content provided."
        else:
            placeholder_p.text = "No content provided."

    # 4. Enable automatic TOC field updates on opening
    enable_update_fields_on_open(doc)

    # 5. Clean up page breaks and layout validation
    # Prevent orphan headings on all Heading paragraphs
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith('Heading'):
            p.paragraph_format.keep_with_next = True

    # Save final dossier
    print(f"Saving compiled dossier to {output_path}...")
    doc.save(output_path)
    
    # Clean up temp folder
    try:
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
    except Exception:
        pass

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python dossier_builder.py <config.json> <output.docx>")
        sys.exit(1)
        
    config_file = sys.argv[1]
    output_file = sys.argv[2]
    compile_dossier(config_file, output_file)
