import docx
from copy import deepcopy

# Create master
d_master = docx.Document()
d_master.add_heading('Master Heading 1', level=1)
p_target = d_master.add_paragraph('Insert after me.')
d_master.add_heading('Master Heading 2', level=1)
d_master.add_paragraph('Tail paragraph.')
d_master.save('test_master.docx')

# Create insert
d_insert = docx.Document()
d_insert.add_heading('Inserted Heading', level=2)
d_insert.add_paragraph('Inserted paragraph.')
d_insert.save('test_insert.docx')

# Perform insertion hack
d_master = docx.Document('test_master.docx')
d_insert = docx.Document('test_insert.docx')

body = d_master.element.body
# Find target paragraph
target_idx = -1
for i, el in enumerate(body):
    if el.tag.endswith('p') and 'Insert after me' in ''.join(el.itertext()):
        target_idx = i
        break

if target_idx != -1:
    tail_elements = body[target_idx+1:-1] # exclude the final sectPr
    final_sectPr = body[-1] if body[-1].tag.endswith('sectPr') else None
    
    # Remove tail elements
    for el in tail_elements:
        body.remove(el)
    if final_sectPr is not None:
        body.remove(final_sectPr)
        
    # Append using manual copy
    elements = list(d_insert.element.body)
    _WNS_local = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    if elements and elements[-1].tag == f"{_WNS_local}sectPr":
        elements = elements[:-1]
    if elements and elements[-1].tag.endswith('sectPr'):
        elements = elements[:-1]
    for element in elements:
        body.append(deepcopy(element))
    
    # Re-append tail elements
    for el in tail_elements:
        body.append(el)
    if final_sectPr is not None:
        body.append(final_sectPr)

d_master.save('test_output.docx')
