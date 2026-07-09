import os
import sys
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

# Add backend directory to path
sys.path.append(str(Path(__file__).parent.resolve()))
import template_components

def build_templates():
    templates_dir = Path(__file__).parent / 'templates'
    templates_dir.mkdir(exist_ok=True)
    print(f"Generating high-fidelity templates in: {templates_dir}")

    # 1. High-Fidelity Cover Page (extracted from amcp_master_template.docx)
    print("Extracting cover.docx...")
    doc_cover = Document(str(Path(__file__).parent / 'amcp_master_template.docx'))
    
    # Remove all sections after cover page
    paragraphs_to_remove = []
    found_editorial = False
    for p in doc_cover.paragraphs:
        if "EDITORIAL STAFF" in p.text or found_editorial:
            found_editorial = True
            paragraphs_to_remove.append(p)
    for p in paragraphs_to_remove:
        p_el = p._element
        p_el.getparent().remove(p_el)
        
    for t in list(doc_cover.tables)[1:]:
        t_el = t._element
        t_el.getparent().remove(t_el)
        
    t0 = doc_cover.tables[0]
    cell = t0.rows[2].cells[1]
    
    # Clear copyright paragraph and add placeholders
    p_cpy = cell.paragraphs[3]
    p_cpy.text = ""
    
    p_brand = cell.add_paragraph()
    p_brand.paragraph_format.space_before = Pt(36)
    p_brand.paragraph_format.space_after = Pt(2)
    r_bl = p_brand.add_run("Product: ")
    r_bl.font.name = 'Asap'
    r_bl.font.bold = True
    r_bl.font.size = Pt(11)
    r_bl.font.color.rgb = template_components.COLOR_JMCP_BLUE
    r_bv = p_brand.add_run("[BRAND_NAME] ([GENERIC_NAME])")
    r_bv.font.name = 'Asap'
    r_bv.font.size = Pt(11)
    r_bv.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_mfg = cell.add_paragraph()
    p_mfg.paragraph_format.space_after = Pt(2)
    r_ml = p_mfg.add_run("Manufacturer: ")
    r_ml.font.name = 'Asap'
    r_ml.font.bold = True
    r_ml.font.size = Pt(11)
    r_ml.font.color.rgb = template_components.COLOR_JMCP_BLUE
    r_mv = p_mfg.add_run("[MANUFACTURER]")
    r_mv.font.name = 'Asap'
    r_mv.font.size = Pt(11)
    r_mv.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_ver = cell.add_paragraph()
    p_ver.paragraph_format.space_after = Pt(2)
    r_vl = p_ver.add_run("Version: ")
    r_vl.font.name = 'Asap'
    r_vl.font.bold = True
    r_vl.font.size = Pt(11)
    r_vl.font.color.rgb = template_components.COLOR_JMCP_BLUE
    r_vv = p_ver.add_run("[VERSION_NUMBER]")
    r_vv.font.name = 'Asap'
    r_vv.font.size = Pt(11)
    r_vv.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_dt = cell.add_paragraph()
    p_dt.paragraph_format.space_after = Pt(36)
    r_dl = p_dt.add_run("Date: ")
    r_dl.font.name = 'Asap'
    r_dl.font.bold = True
    r_dl.font.size = Pt(11)
    r_dl.font.color.rgb = template_components.COLOR_JMCP_BLUE
    r_dv = p_dt.add_run("[COMPILATION_DATE]")
    r_dv.font.name = 'Asap'
    r_dv.font.size = Pt(11)
    r_dv.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_cpy_new = cell.add_paragraph()
    r_cpy = p_cpy_new.add_run("Copyright © 2024 AMCP")
    r_cpy.font.name = 'Asap'
    r_cpy.font.size = Pt(7.5)
    r_cpy.font.color.rgb = template_components.COLOR_SLATE_GRAY
    doc_cover.save(str(templates_dir / 'cover.docx'))

    # 2. Table of Contents Template
    print("Generating toc.docx...")
    doc_toc = Document(str(Path(__file__).parent / 'amcp_template.docx'))
    p_remove = list(doc_toc.paragraphs)[4:]
    for p in p_remove:
        p_el = p._element
        p_el.getparent().remove(p_el)
    p0 = doc_toc.paragraphs[0]
    p0_el = p0._element
    p0_el.getparent().remove(p0_el)
    t0_el = doc_toc.tables[0]._element
    t0_el.getparent().remove(t0_el)
    doc_toc.save(str(templates_dir / 'toc.docx'))

    # 3. Standard Content Template
    print("Generating standard_content.docx...")
    doc_std = Document(str(Path(__file__).parent / 'amcp_template.docx'))
    p_remove = list(doc_std.paragraphs)[:4]
    for p in p_remove:
        p_el = p._element
        p_el.getparent().remove(p_el)
    t0_el = doc_std.tables[0]._element
    t0_el.getparent().remove(t0_el)
    for p in list(doc_std.paragraphs)[1:]:
        p_el = p._element
        p_el.getparent().remove(p_el)
    doc_std.paragraphs[0].text = "[CONTENT_PLACEHOLDER]"
    doc_std.save(str(templates_dir / 'standard_content.docx'))

    # 4. Section Divider Template
    print("Generating section_header.docx...")
    doc_sec = Document()
    template_components.apply_global_theme(doc_sec)
    section = doc_sec.sections[0]
    template_components.configure_section_layout(section, "content")
    template_components.setup_headers_footers(section)
    p_heading = doc_sec.add_paragraph()
    template_components.create_section_heading_component(p_heading, "[SECTION_NUMBER]", "[SECTION_TITLE]")
    doc_sec.save(str(templates_dir / 'section_header.docx'))

    # 5. Table Template
    print("Generating table.docx...")
    doc_tbl = Document()
    template_components.apply_global_theme(doc_tbl)
    table = doc_tbl.add_table(3, 3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    template_components.clear_table_borders(table)
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for b_name in ['top', 'bottom', 'insideH']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), template_components.HEX_SLATE_GRAY)
        tblBorders.append(b)
    for b_name in ['left', 'right', 'insideV']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    hdr_row = table.rows[0]
    for cell in hdr_row.cells:
        template_components.set_cell_shading(cell, template_components.HEX_JMCP_BLUE)
        template_components.set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run("Header Cell")
        r.font.name = 'Asap'
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for idx, row in enumerate(table.rows[1:], start=1):
        shading_color = template_components.HEX_VERY_LIGHT_GRAY if idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            if shading_color != "FFFFFF":
                template_components.set_cell_shading(cell, shading_color)
            template_components.set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run("Data Cell")
            r.font.name = 'Lora'
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0, 0, 0)
    doc_tbl.save(str(templates_dir / 'table.docx'))

    # 6. Build the static, unified master templates format-a and format-b
    build_unified_master_template("format-a", templates_dir / 'amcp_template_format_a.docx')
    build_unified_master_template("format-b", templates_dir / 'amcp_template_format_b.docx')

    print("All master templates generated successfully!")

# Combines Cover, TOC, and content placeholders into a single static file
def build_unified_master_template(dossier_type, output_path):
    print(f"Building unified template for {dossier_type}...")
    templates_dir = Path(__file__).parent / 'templates'
    master = Document(str(templates_dir / 'cover.docx'))
    composer = Composer(master)
    
    # Append TOC
    doc_toc = Document(str(templates_dir / 'toc.docx'))
    composer.append(doc_toc)
    
    # Map section details
    format_letter = 'B' if dossier_type == "format-b" else 'A'
    prod_type_text = 'Approved Product' if dossier_type == "format-b" else 'Unapproved Product / Unapproved Use'
    
    section_titles = [
        ('highlights' if format_letter == 'A' else 'exec_summary', f"1.0{format_letter}", f"Executive Summary: Clinical and Economic Value of the {prod_type_text}" if format_letter == 'B' else f"Highlights and Overview of the {prod_type_text}"),
        ('product_info', f"2.0{format_letter}", "Product Information and Disease Description"),
        ('clinical', f"3.0{format_letter}", "Clinical Evidence"),
        ('economic', f"4.0{format_letter}", "Economic Value and Modeling Report"),
        ('supporting', f"5.0{format_letter}", "Additional Supporting Evidence"),
        ('references', f"6.0{format_letter}", "Dossier Appendices and References")
    ]
    
    for key, num, title in section_titles:
        # Append Section Heading
        doc_h = Document(str(templates_dir / 'section_header.docx'))
        p = doc_h.paragraphs[0]
        # set text inside template heading
        p.text = ""
        cyan_run = p.add_run(f"SECTION {num} \n")
        cyan_run.font.name = 'Asap'
        cyan_run.font.size = Pt(10)
        cyan_run.font.bold = True
        cyan_run.font.color.rgb = RGBColor(0, 126, 197)
        plum_run = p.add_run(title)
        plum_run.font.name = 'Asap'
        plum_run.font.size = Pt(16)
        plum_run.font.bold = True
        plum_run.font.color.rgb = RGBColor(108, 16, 91)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        composer.append(doc_h)
        
        # Append Content Placeholder
        doc_c = Document(str(templates_dir / 'standard_content.docx'))
        p_placeholder = doc_c.paragraphs[0]
        p_placeholder.text = f"[INSERT_{key.upper()}]"
        composer.append(doc_c)
        
    # Standardize margins and geometry
    for idx, section in enumerate(master.sections):
        if idx == 0:
            continue
        section.page_width = Inches(8.25)
        section.page_height = Inches(10.875)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        section.even_page_header.is_linked_to_previous = True
        section.even_page_footer.is_linked_to_previous = True
        
        if idx >= 2:
            sectPr = section._sectPr
            cols = sectPr.xpath('w:cols')
            if cols:
                cols[0].set(qn('w:num'), '2')
                cols[0].set(qn('w:space'), '360')
            else:
                col_el = OxmlElement('w:cols')
                col_el.set(qn('w:num'), '2')
                col_el.set(qn('w:space'), '360')
                sectPr.append(col_el)
        else:
            sectPr = section._sectPr
            cols = sectPr.xpath('w:cols')
            if cols:
                cols[0].set(qn('w:num'), '1')
                
    master.save(str(output_path))
    print(f"Successfully generated master template: {output_path}")

if __name__ == '__main__':
    build_templates()
